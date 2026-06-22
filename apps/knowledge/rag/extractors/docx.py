"""DOCX 文件提取器。"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


async def extract_docx(file_path: str) -> str:
    """从 docx 文件提取 markdown 文本。"""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("需要安装 python-docx: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"文件不存在: {file_path}")

    doc = Document(file_path)

    md_parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""
        if "heading" in style_name:
            level = 1
            for i in range(1, 7):
                if f"heading {i}" in style_name:
                    level = i
                    break
            md_parts.append(f"{'#' * level} {text}")
        else:
            md_parts.append(text)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if rows:
            if len(rows) > 1:
                separator = "| " + " | ".join(["---"] * len(table.columns)) + " |"
                rows.insert(1, separator)
            md_parts.append("\n".join(rows))

    result = "\n\n".join(md_parts)
    if not result.strip():
        raise ValueError("docx 文件内容为空或无法提取文本")

    return result
